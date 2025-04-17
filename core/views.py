from django.shortcuts import render, redirect , get_object_or_404
from django.contrib.auth import authenticate, login ,logout
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import Invoice, Customer, Design, Stone, Design, InventoryTransaction , Expense , FinancialReport , CustomerPayment
from django.utils import timezone
from django.db.models import Q , Sum , F
from django.contrib.auth.models import User
from decimal import Decimal
from django.utils.dateparse import parse_date
from collections import defaultdict
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from django.utils.text import slugify
from datetime import date
from persiantools.jdatetime import JalaliDate
import jdatetime
from django.http import FileResponse, Http404
from django.conf import settings
import os
from django.core.paginator import Paginator


# Create your views here.

def convert_persian_numbers(s):
    persian_numbers = '۰۱۲۳۴۵۶۷۸۹'
    english_numbers = '0123456789'
    return s.translate(str.maketrans(persian_numbers, english_numbers))


def login_view(request):
    # اگر کاربر وارد شده بود، هدایت به داشبورد
    if request.user.is_authenticated:
        return redirect('home')

    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')

        # بررسی صحت اطلاعات
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('home')  # هدایت به داشبورد بعد از ورود موفق
        else:
            messages.error(request, "نام کاربری یا رمز عبور اشتباه است.")

    return render(request, 'login.html')

def logout_view(request):
    logout(request)
    return redirect('login')

# صفحه اصلی home 
@login_required
def home_view(request):
    latest_invoices = Invoice.objects.order_by('-date')[:5]  # آخرین ۵ فاکتور
    stones = Stone.objects.all()
    return render(request, 'home.html', {
        'latest_invoices': latest_invoices ,
        'stones': stones, 
    })

# صفحه مشتری ها
@login_required
def customer_list(request):
    query = request.GET.get('q')
    if query:
        customers = Customer.objects.filter(name__icontains=query)
    else:
        customers = Customer.objects.all().order_by('name')

    # افزودن مشتری
    if request.method == 'POST' and 'add_customer' in request.POST:
        name = request.POST.get('name')
        phone = request.POST.get('phone')
        address = request.POST.get('address')
        Customer.objects.create(name=name, phone=phone, address=address)
        messages.success(request, "مشتری جدید با موفقیت اضافه شد.")
        return redirect('customers')

    # ویرایش مشتری
    if request.method == 'POST' and 'edit_customer' in request.POST:
        customer_id = request.POST.get('customer_id')
        customer = get_object_or_404(Customer, id=customer_id)
        customer.name = request.POST.get('name')
        customer.phone = request.POST.get('phone')
        customer.address = request.POST.get('address')
        customer.save()
        messages.success(request, "اطلاعات مشتری ویرایش شد.")
        return redirect('customers')

    # ثبت پرداخت
    if request.method == 'POST' and 'add_payment' in request.POST:
        customer_id = request.POST.get('customer_id')
        user_id = request.POST.get('received_by')  # ✅ کاربری که پول را دریافت کرده
        payment = int(request.POST.get('payment_amount').replace(',', ''))
        payment_date = request.POST.get('date').replace('/', '-')

        customer = get_object_or_404(Customer, id=customer_id)
        receiver = get_object_or_404(User, id=user_id)

        # محاسبه بدهی فعلی
        total_invoice = Invoice.objects.filter(customer=customer).aggregate(Sum('total_price'))['total_price__sum'] or 0
        total_paid = Invoice.objects.filter(customer=customer).aggregate(Sum('amount_paid'))['amount_paid__sum'] or 0
        current_debt = total_invoice - total_paid

        if payment > current_debt:
            messages.error(request, f"مبلغ وارد شده بیشتر از بدهی مشتری ({current_debt:,} تومان) است.")
            return redirect('customers')

        # تسویه فاکتورهای بدهکار
        remaining = payment
        unpaid_invoices = Invoice.objects.filter(customer=customer, remaining_debt__gt=0).order_by('date')
        for invoice in unpaid_invoices:
            if remaining <= 0:
                break

            if remaining >= invoice.remaining_debt:
                paid_amount = invoice.remaining_debt
                remaining -= invoice.remaining_debt
                invoice.amount_paid += paid_amount
                invoice.remaining_debt = 0
            else:
                paid_amount = remaining
                invoice.amount_paid += paid_amount
                invoice.remaining_debt -= paid_amount
                remaining = 0
            invoice.save()

            # ✅ ثبت پرداخت با اتصال به فاکتور
            CustomerPayment.objects.create(
                customer=customer,
                amount=paid_amount,
                date=payment_date,
                user=request.user,
                received_by=receiver
            )

        messages.success(request, f"پرداخت {payment:,} تومان برای مشتری {customer.name} ثبت شد.")
        return redirect('customers')

    # ساخت customer_data
    customer_data = []
    invoices_by_customer = {}

    for customer in customers:
        invoices = Invoice.objects.filter(customer=customer).order_by('-date')
        total_invoice = invoices.aggregate(Sum('total_price'))['total_price__sum'] or 0
        total_paid = invoices.aggregate(Sum('amount_paid'))['amount_paid__sum'] or 0
        total_debt = total_invoice - total_paid

        customer_data.append({
            'customer': customer,
            'total_invoice': total_invoice,
            'total_paid': total_paid,
            'total_debt': total_debt,
        })

        invoices_by_customer[customer.id] = invoices

    return render(request, 'customers/list.html', {
        'customer_data': customer_data,
        'customers': customers,
        'invoices_by_customer': invoices_by_customer,
        'users': User.objects.all(),  # 👈 نمایش کاربران برای انتخاب در فرم
    })

    
# فاکتورهای فروش
@login_required
def invoice_list(request):
    query = request.GET.get('q')
    invoices = Invoice.objects.select_related('customer', 'design', 'stone').order_by('-date')
    if query:
        invoices = invoices.filter(
            Q(customer__name__icontains=query) |
            Q(invoice_number__icontains=query)
        )
        
          
    # 🔁✏️ ویرایش فاکتور موجود
    if request.method == 'POST' and 'edit_invoice' in request.POST:
        invoice = get_object_or_404(Invoice, id=request.POST.get('invoice_id'))
        
        invoice.date = request.POST.get('date').replace('/', '-')
        invoice.customer_id = request.POST.get('customer')
        invoice.design_id = request.POST.get('design')
        invoice.stone_id = request.POST.get('stone')
        invoice.design_price_per_piece = int(request.POST.get('design_price_per_piece').replace(',', '') or 0)
        invoice.quantity = int(request.POST.get('quantity') or 0)
        invoice.discount = int(request.POST.get('discount').replace(',', '') or 0)
        invoice.total_price = (invoice.design_price_per_piece * invoice.quantity) - invoice.discount
        invoice.amount_paid = int(request.POST.get('amount_paid').replace(',', '') or 0)
        invoice.remaining_debt = invoice.total_price - invoice.amount_paid
        invoice.payment_type = request.POST.get('payment_type')
        invoice.check_due_date = request.POST.get('check_due_date') or None
        invoice.note = request.POST.get('note') or ''
        invoice.issuer_id = request.POST.get('issuer')
        
        invoice.save()
        messages.success(request, f"فاکتور {invoice.invoice_number} با موفقیت ویرایش شد.")
        return redirect('invoices')



    # افزودن فاکتور جدید
    if request.method == 'POST' and 'add_invoice' in request.POST:
        # بررسی اینکه آیا مشتری جدید وارد شده یا مشتری انتخابی استفاده شده
        customer = None
        new_name = request.POST.get('new_customer_name', '').strip()
        new_phone = request.POST.get('new_customer_phone', '').strip()
        new_address = request.POST.get('new_customer_address', '').strip()

        if new_name:
            customer = Customer.objects.create(
                name=new_name,
                phone=new_phone,
                address=new_address
            )
        else:
            customer_id = request.POST.get('customer')
            if not customer_id:
                messages.error(request, "لطفاً یک مشتری انتخاب کنید یا مشتری جدید وارد نمایید.")
                return redirect('invoices')
            customer = get_object_or_404(Customer, id=customer_id)


        # 2. دریافت اطلاعات فاکتور
        design_id = request.POST.get('design')
        stone_id = request.POST.get('stone')
        design_price_per_piece = int(request.POST.get('design_price_per_piece').replace(',', '') or '0')
        quantity = int(request.POST.get('quantity') or '0')
        discount = int(request.POST.get('discount').replace(',', '') or '0')
        amount_paid = int(request.POST.get('amount_paid').replace(',', '') or '0')
        payment_type = request.POST.get('payment_type')
        raw_due = request.POST.get('check_due_date')
        check_due_date = raw_due.replace('/', '-') if raw_due else None
        note = request.POST.get('note') or ''
        invoice_date = request.POST.get('date').replace('/', '-')
        issuer_id = request.POST.get('issuer')

        # 3. ساخت شماره فاکتور متوالی ۸ رقمی
        last_invoice = Invoice.objects.order_by('-id').first()
        if last_invoice and last_invoice.invoice_number.startswith("INV"):
            try:
                last_number = int(last_invoice.invoice_number.replace("INV", ""))
            except:
                last_number = 1
        else:
            last_number = 1
        invoice_number = f"INV{last_number + 1}"

        # 4. محاسبه مبلغ کل و بدهی
        total_price = (design_price_per_piece * quantity) - discount
        remaining_debt = total_price - amount_paid

        # 5. ساخت فاکتور
        invoice = Invoice.objects.create(
            date=invoice_date,
            invoice_number=invoice_number,
            customer=customer,
            design_id=design_id,
            stone_id=stone_id,
            design_price_per_piece=design_price_per_piece,
            quantity=quantity,
            discount=discount,
            total_price=total_price,
            issuer_id=issuer_id,
            payment_type=payment_type,
            check_due_date=check_due_date,
            amount_paid=amount_paid,
            remaining_debt=remaining_debt,
            note=note
        )

        # ✅ 6. اگر مبلغی پرداخت شده، ثبت CustomerPayment
        if amount_paid > 0:
            CustomerPayment.objects.create(
                customer=customer,
                invoice=invoice,  # ✅ فاکتور جدیدی که همین الان ساختیم
                amount=amount_paid,
                date=invoice_date,
                user=request.user,
                received_by=User.objects.get(id=issuer_id)
            )

        messages.success(request, "فاکتور با موفقیت ثبت شد.")
        return redirect('invoices')
    
    
    # بعد از دریافت لیست کامل:
    paginator = Paginator(invoices, 20)  # نمایش ۲۰ فاکتور در هر صفحه
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # ارسال به قالب:
    return render(request, 'invoices/list.html', {
        'page_obj': page_obj,
        'invoices': page_obj.object_list,
        'customers': Customer.objects.all(),
        'designs': Design.objects.all(),
        'stones': Stone.objects.all(),
        'users': User.objects.all(),
    })


@login_required
def inventory_list(request):
    # ثبت سنگ جدید
    if request.method == 'POST' and 'add_stone' in request.POST:
        name = request.POST.get('name')
        price_per_box_usd = request.POST.get('price_per_box_usd')
        Stone.objects.create(name=name, price_per_box_usd=price_per_box_usd)
        messages.success(request, "سنگ جدید ثبت شد.")
        return redirect('inventory')

    # ثبت طرح جدید
    if request.method == 'POST' and 'add_design' in request.POST:
        title = request.POST.get('title')
        Design.objects.create(title=title)
        messages.success(request, "طرح جدید ثبت شد.")
        return redirect('inventory')

    # ثبت تراکنش سنگ
    if request.method == 'POST' and 'add_transaction' in request.POST:
        stone_id = request.POST.get('stone')
        quantity = Decimal(request.POST.get('quantity_box'))
        date = request.POST.get('date').replace('/', '-')
        type = request.POST.get('type')

        # ثبت تراکنش
        transaction = InventoryTransaction.objects.create(
            stone_id=stone_id,
            quantity_box=quantity,
            date=date,
            type=type
        )

        # به‌روزرسانی موجودی سنگ
        stone = transaction.stone
        if type == 'buy':
            stone.stock_box += quantity
        elif type == 'consume':
            stone.stock_box -= quantity
        stone.save()

        messages.success(request, "تراکنش با موفقیت ثبت شد.")
        return redirect('inventory')

    stones = Stone.objects.all().order_by('name')
    designs = Design.objects.all().order_by('title')
    transactions = InventoryTransaction.objects.select_related('stone').order_by('-date')[:10]  # آخرین ۱۰ تراکنش

    return render(request, 'inventory/list.html', {
        'stones': stones,
        'designs': designs,
        'transactions': transactions,
    })

@login_required
def expense_list(request):
    expenses = Expense.objects.select_related('user').order_by('-date')

    # فیلتر براساس تاریخ
    start = request.GET.get('start_date')
    end = request.GET.get('end_date')
    if start and end:
        start_date = parse_date(start.replace('/', '-'))
        end_date = parse_date(end.replace('/', '-'))
        expenses = expenses.filter(date__range=(start_date, end_date))

    # افزودن هزینه جدید
    if request.method == 'POST' and 'add_expense' in request.POST:
        Expense.objects.create(
            date=request.POST.get('date').replace('/', '-'),
            user_id=request.POST.get('user'),
            amount = int(request.POST.get('amount').replace(',', '')),
            note=request.POST.get('note')[0:50]
        )

        messages.success(request, "هزینه جدید با موفقیت ثبت شد.")
        return redirect('expenses')

    return render(request, 'expenses/list.html', {
        'expenses': expenses,
        'users': User.objects.all(),
    })

@login_required
def financial_report(request):
    report = None
    invoices = []
    expenses = []
    stone_cost = Decimal(0)
    customers_data = {}
    users_data = {}
    total_received = 0
    total_unpaid = 0

    if request.method == 'POST':
        start_raw = request.POST.get('start_date').replace('/', '-')
        end_raw = request.POST.get('end_date').replace('/', '-')

        try:
            start_jdate = jdatetime.datetime.strptime(start_raw, "%Y-%m-%d")
            end_jdate = jdatetime.datetime.strptime(end_raw, "%Y-%m-%d")
            start_date = start_jdate.togregorian().date()
            end_date = end_jdate.togregorian().date()
        except Exception as e:
            messages.error(request, "خطا در تبدیل تاریخ: " + str(e))
            return redirect('financial_report')

        if start_date > end_date:
            start_date, end_date = end_date, start_date

        usd_rate_raw = request.POST.get('usd_rate')
        if not usd_rate_raw:
            messages.error(request, "نرخ دلار وارد نشده است.")
            return redirect('financial_report')
        usd_rate = Decimal(usd_rate_raw.replace(',', ''))

        # 📦 فاکتورها
        invoices = Invoice.objects.select_related('issuer', 'customer').filter(date__range=(start_date, end_date)).order_by('-date')
        total_pieces_sold = sum(i.quantity for i in invoices)
        total_sales = sum(i.total_price for i in invoices)

        # 💸 هزینه‌های دیگر
        expenses = Expense.objects.select_related('user').filter(date__range=(start_date, end_date)).order_by('-date')
        total_expenses = sum(e.amount for e in expenses)

        # 💎 هزینه سنگ‌ها
        transactions = InventoryTransaction.objects.select_related('stone').filter(
            type='consume', date__range=(start_date, end_date)
        ).order_by('-date')
        for t in transactions:
            stone_cost += t.quantity_box * t.stone.price_per_box_usd * usd_rate

        # 💰 سود خالص
        net_profit = total_sales - stone_cost - total_expenses

        # 📥 عملکرد کاربران
        users_data = {}
        for user in User.objects.all():
            users_data[user.id] = {
                'username': user.username,
                'invoice_received': 0,
                'manual_received': 0,
                'total_debt': 0,
                'check_count': 0,
                'checks': [],
            }


        # ✅ جمع‌آوری داده‌های اولیه فاکتورها و مشتری‌ها
        for invoice in invoices:
            user = invoice.issuer
            data = users_data[user.id]
            data['username'] = user.username
            data['total_debt'] += invoice.remaining_debt

            if invoice.payment_type == 'check':
                data['check_count'] += 1
                data['checks'].append({
                    'date': invoice.check_due_date,
                    'note': invoice.note or '—',
                })

            # مشتریان دارای فاکتور
            c = invoice.customer
            if c.id not in customers_data:
                customers_data[c.id] = {
                    'name': c.name,
                    'invoices': [],
                    'total_quantity': 0,
                    'total_paid': 0,
                    'total_debt': 0,
                }
            customers_data[c.id]['invoices'].append(invoice)
            customers_data[c.id]['total_quantity'] += invoice.quantity
            customers_data[c.id]['total_paid'] += invoice.amount_paid
            customers_data[c.id]['total_debt'] += invoice.remaining_debt

        from core.models import CustomerPayment  # در صورت نیاز ایمپورت کن

        # ✅ پرداخت‌های CustomerPayment در بازه
        payments = CustomerPayment.objects.select_related('invoice', 'received_by', 'invoice__issuer').filter(date__range=(start_date, end_date))

        for p in payments:
            if p.invoice and p.received_by:
                if p.invoice.issuer == p.received_by:
                    users_data[p.received_by.id]['invoice_received'] += p.amount
                else:
                    users_data[p.received_by.id]['manual_received'] += p.amount
            elif p.received_by:
                users_data[p.received_by.id]['manual_received'] += p.amount


        # محاسبه دریافتی نهایی هر کاربر
        for u in users_data.values():
            u['net_received'] = u['invoice_received'] + u['manual_received']

        # محاسبه کل دریافتی و بدهی
        total_received = sum(u['net_received'] for u in users_data.values())
        total_unpaid = sum(u['total_debt'] for u in users_data.values())

        users_data = dict(users_data)

        # ذخیره گزارش
        report = FinancialReport.objects.create(
            start_date=start_date,
            end_date=end_date,
            usd_rate=usd_rate,
            total_pieces_sold=total_pieces_sold,
            total_sales_amount=total_sales,
            total_stone_cost=int(stone_cost),
            total_expenses=total_expenses,
            net_profit=int(net_profit),
            created_by=request.user,
        )

        messages.success(request, "گزارش مالی با موفقیت ذخیره شد.")

    return render(request, 'financial/list.html', {
        'report': report,
        'invoices': invoices,
        'expenses': expenses,
        'customers_data': customers_data,
        'users_data': users_data,
        'total_received': total_received,
        'total_unpaid': total_unpaid,
    })



def jalali_to_gregorian(date_str):
    try:
        y, m, d = map(int, date_str.split('/'))
        return JalaliDate(y, m, d).to_gregorian()
    except:
        return None


@login_required
def user_report(request):
    start = request.GET.get('start_date')
    end = request.GET.get('end_date')

    start_date = jalali_to_gregorian(start) if start else None
    end_date = jalali_to_gregorian(end) if end else None

    # اگر تاریخ‌ها جابه‌جا بودن، درستش کن
    if start_date and end_date and start_date > end_date:
        start_date, end_date = end_date, start_date

    users = User.objects.all()
    user_data = []

    for user in users:
        invoices = Invoice.objects.filter(issuer=user)
        expenses = Expense.objects.filter(user=user)
        payments = CustomerPayment.objects.filter(received_by=user)

        if start_date and end_date:
            invoices = invoices.filter(date__range=(start_date, end_date))
            expenses = expenses.filter(date__range=(start_date, end_date))
            payments = payments.filter(date__range=(start_date, end_date))

        # محاسبه دریافتی واقعی = پرداخت‌های مرتبط با فاکتور + پرداخت‌های مستقیم
        total_received = 0
        for p in payments:
            if p.invoice and p.invoice.issuer == p.received_by:
                total_received += p.amount
            elif not p.invoice:
                total_received += p.amount

        total_expenses = expenses.aggregate(Sum('amount'))['amount__sum'] or 0

        checks = invoices.filter(payment_type='check')
        check_count = checks.count()
        check_details = [{'date': i.check_due_date, 'note': i.note or '—'} for i in checks]

        user_data.append({
            'user': user,
            'net_received': total_received,
            'total_expenses': total_expenses,
            'check_count': check_count,
            'checks': check_details,
            'invoices': invoices.order_by('-date'),
            'expenses': expenses.order_by('-date'),
        })

    return render(request, 'users/report.html', {
        'user_data': user_data,
        'start_date': start,
        'end_date': end,
    })


@login_required
def export_financial_report_docx(request):
    report = FinancialReport.objects.order_by('-created_at').first()
    if not report:
        messages.error(request, "هیچ گزارشی برای چاپ وجود ندارد.")
        return redirect('financial')

    # داده‌ها
    start_date, end_date = report.start_date, report.end_date
    usd_rate = report.usd_rate

    invoices = Invoice.objects.select_related('customer', 'issuer').filter(date__range=(start_date, end_date))
    expenses = Expense.objects.select_related('user').filter(date__range=(start_date, end_date))
    payments = CustomerPayment.objects.select_related('invoice', 'received_by').filter(date__range=(start_date, end_date))

    # ساخت سند
    doc = Document()
    doc.add_heading('📄 گزارش مالی سیستم کارگاه', 0)

    doc.add_paragraph(f'🗓️ بازه گزارش: {start_date.strftime("%Y/%m/%d")} تا {end_date.strftime("%Y/%m/%d")}')
    doc.add_paragraph(f'💵 نرخ دلار: {usd_rate:,} تومان')
    doc.add_paragraph(f'👤 ثبت‌کننده گزارش: {report.created_by.username}')
    doc.add_paragraph('')

    # خلاصه مالی
    doc.add_heading('📈 خلاصه گزارش مالی', level=1)
    doc.add_paragraph(f'• مجموع قواره‌های فروش‌رفته: {report.total_pieces_sold:,}')
    doc.add_paragraph(f'• جمع فروش: {report.total_sales_amount:,} تومان')
    doc.add_paragraph(f'• جمع هزینه سنگ: {report.total_stone_cost:,} تومان')
    doc.add_paragraph(f'• جمع سایر هزینه‌ها: {report.total_expenses:,} تومان')
    doc.add_paragraph(f'• 💰 سود خالص: {report.net_profit:,} تومان')
    doc.add_paragraph('')

    # لیست مشتریان
    doc.add_heading('🧾 مشتریان دارای فاکتور در بازه', level=1)
    customers = {}
    for i in invoices:
        c = i.customer
        if c.id not in customers:
            customers[c.id] = {'name': c.name, 'q': 0, 'paid': 0, 'debt': 0}
        customers[c.id]['q'] += i.quantity
        customers[c.id]['paid'] += i.amount_paid
        customers[c.id]['debt'] += i.remaining_debt

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'نام مشتری'
    hdr[1].text = 'مجموع قواره'
    hdr[2].text = 'پرداختی'
    hdr[3].text = 'بدهی'

    for c in customers.values():
        row = table.add_row().cells
        row[0].text = c['name']
        row[1].text = f"{c['q']:,}"
        row[2].text = f"{c['paid']:,}"
        row[3].text = f"{c['debt']:,}"
    doc.add_paragraph('')

    # لیست هزینه‌ها
    doc.add_heading('💰 هزینه‌های انجام‌شده', level=1)
    exp_table = doc.add_table(rows=1, cols=4)
    exp_table.style = 'Table Grid'
    h = exp_table.rows[0].cells
    h[0].text = 'تاریخ'
    h[1].text = 'کاربر'
    h[2].text = 'مبلغ'
    h[3].text = 'توضیح'

    for e in expenses:
        row = exp_table.add_row().cells
        row[0].text = str(e.date)
        row[1].text = e.user.username
        row[2].text = f"{e.amount:,}"
        row[3].text = e.note or '—'
    doc.add_paragraph('')

    # 👥 عملکرد کاربران (اصلاح‌شده برای نمایش همه کاربران)
    doc.add_heading('👥 عملکرد کاربران', level=1)

    # آماده‌سازی اولیه
    users = {}
    for user in User.objects.all():
        users[user.id] = {
            'username': user.username,
            'sales': 0,
            'debt': 0,
            'received': 0,
            'checks': []
        }

    # جمع‌آوری داده‌ها از فاکتورها
    for i in invoices:
        u = i.issuer
        if u:
            users[u.id]['sales'] += i.total_price
            users[u.id]['debt'] += i.remaining_debt
            if i.payment_type == 'check':
                users[u.id]['checks'].append({'date': str(i.check_due_date), 'note': i.note or '—'})

    # جمع‌آوری داده‌ها از پرداخت‌ها
    payments = CustomerPayment.objects.select_related('received_by', 'invoice__issuer').filter(
        date__range=(report.start_date, report.end_date)
    )
    for p in payments:
        if p.received_by:
            users[p.received_by.id]['received'] += p.amount

    # نمایش جدول کاربران
    utable = doc.add_table(rows=1, cols=5)
    utable.style = 'Table Grid'
    h = utable.rows[0].cells
    h[0].text = 'کاربر'
    h[1].text = 'دریافتی واقعی'
    h[2].text = 'بدهی'
    h[3].text = 'فروش کل'
    h[4].text = 'تعداد چک'

    for u in users.values():
        row = utable.add_row().cells
        row[0].text = u['username']
        row[1].text = f"{u['received']:,}"
        row[2].text = f"{u['debt']:,}"
        row[3].text = f"{u['sales']:,}"
        row[4].text = str(len(u['checks']))

        if u['checks']:
            doc.add_paragraph(f"📋 چک‌های {u['username']}:", style='List Bullet')
            for ch in u['checks']:
                doc.add_paragraph(f"• تاریخ: {ch['date']} - توضیح: {ch['note']}", style='List Bullet 2')

    # 📄 تولید و ارسال فایل
    filename = f"گزارش_مالی_{slugify(str(start_date))}_تا_{slugify(str(end_date))}.docx"
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)
    return response



@login_required
def download_backup(request):
    db_path = settings.DATABASES['default']['NAME']

    if os.path.exists(db_path):
        return FileResponse(open(db_path, 'rb'), as_attachment=True, filename='backup.sqlite3')
    else:
        raise Http404("فایل دیتابیس پیدا نشد.")
